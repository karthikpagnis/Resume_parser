"""Utility functions for Resume Parser"""

import re
import json
from pathlib import Path
from typing import Dict, List, Tuple
import subprocess
import platform

import nltk
from nltk.tokenize import sent_tokenize
import PyPDF2
from docx import Document

from config import EMAIL_PATTERN, PHONE_PATTERN, LINKEDIN_PATTERN, GITHUB_PATTERN

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')


class FileExtractor:
    """Extract text from various file formats"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF files"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
        return text
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX files"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")
        return text
    
    @staticmethod
    def extract_from_doc(file_path: str) -> str:
        """Extract text from DOC files using system tools"""
        text = ""
        try:
            system = platform.system()
            if system == "Darwin":  # macOS
                cmd = f"textutil -convert txt '{file_path}' -output /tmp/resume_temp.txt && cat /tmp/resume_temp.txt"
                result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
                text = result.stdout
            elif system == "Linux":
                cmd = f"libreoffice --headless --convert-to txt '{file_path}' --outdir /tmp && cat /tmp/resume_temp.txt"
                result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
                text = result.stdout
            elif system == "Windows":
                # Using python-docx to read doc files (it handles both .doc and .docx)
                text = FileExtractor.extract_from_docx(file_path)
        except Exception as e:
            print(f"Error reading DOC {file_path}: {e}")
            # Fallback to docx extraction
            text = FileExtractor.extract_from_docx(file_path)
        return text
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from file based on extension"""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            return FileExtractor.extract_from_pdf(file_path)
        elif extension == '.docx':
            return FileExtractor.extract_from_docx(file_path)
        elif extension in ['.doc', '.docs']:
            return FileExtractor.extract_from_doc(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")


class InformationExtractor:
    """Extract specific information from resume text"""
    
    @staticmethod
    def extract_emails(text: str) -> List[str]:
        """Extract email addresses"""
        return re.findall(EMAIL_PATTERN, text)
    
    @staticmethod
    def extract_phones(text: str) -> List[str]:
        """Extract phone numbers"""
        return re.findall(PHONE_PATTERN, text)
    
    @staticmethod
    def extract_linkedin(text: str) -> List[str]:
        """Extract LinkedIn profile"""
        return re.findall(LINKEDIN_PATTERN, text, re.IGNORECASE)
    
    @staticmethod
    def extract_github(text: str) -> List[str]:
        """Extract GitHub profile"""
        return re.findall(GITHUB_PATTERN, text, re.IGNORECASE)
    
    @staticmethod
    def extract_urls(text: str) -> List[str]:
        """Extract URLs"""
        url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        return re.findall(url_pattern, text)
    
    @staticmethod
    def extract_sentences(text: str) -> List[str]:
        """Extract sentences from text"""
        return sent_tokenize(text)
    
    @staticmethod
    def get_section_content(text: str, section_keywords: List[str]) -> str:
        """Extract content after section headers"""
        # Split by common section headers
        sections = re.split(
            r'(?:^|\n)(?:' + '|'.join(section_keywords) + r')(?:\s*[:\-]?)',
            text,
            flags=re.IGNORECASE | re.MULTILINE
        )
        return sections[1] if len(sections) > 1 else ""


class TextCleaner:
    """Clean and normalize text"""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean resume text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep emails, urls, etc
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text.strip()
    
    @staticmethod
    def normalize_text(text: str) -> str:
        """Normalize text for processing"""
        text = TextCleaner.clean_text(text)
        return text.lower()


def save_json_output(data: Dict, output_path: str) -> None:
    """Save extracted data to JSON file"""
    with open(output_path, 'w') as f:
        json.dump(data, f, indent=2)


def load_json_output(json_path: str) -> Dict:
    """Load JSON output file"""
    with open(json_path, 'r') as f:
        return json.load(f)
